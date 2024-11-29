import shutil
import subprocess
import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from selenium import webdriver
from selenium.webdriver.common.by import By
from time import sleep
from bs4 import BeautifulSoup, Comment
import requests
import xml.etree.ElementTree as ET
import difflib
import re
import json
import argparse

parser = argparse.ArgumentParser()
parser.add_argument("-T", metavar="<target>", action="store", type=str, help="Target IP address", required=True)
parser.add_argument("-L", action="store_true", help="Use to attempt auto logging in. WARNING! CAN TAKE A LONG TIME")
parser.add_argument("-S", action="store_true", help="Use to take screenshots of pages automatically")
parser.add_argument("-A", action="store_true", help="Turn all scripts on. Use if testing on a live machine")
parser.add_argument("-P", action="store_true", help="Create a pdf document of report aswell. REQUIRES LIBRE OFFICE")
parser.add_argument("-O", action="store_true", help="Use provided example scan folder to create report in both docx and pdf format. Used only to demonstrate creation of report with a provided output folder")
args = parser.parse_args()

target = None
loginScript = False
screenshotScript = False
outputOnly = False
pdf = False

if args.T:
    target = args.T

if args.L:
    # Code for switch1
    loginScript = True
    print("Auto login is enabled")

if args.S:
    # Code for switch2
    screenshotScript = True
    print(f"Screen shot taking is on")

if args.A:
    loginScript = True
    screenshotScript = True
    print(f"All scripts are on")

if args.P:
    pdf = True
    print(f"Will create pdf report alongside docx report")

if args.O:
    outputOnly = True
    pdf = True
    print(f"Demonstrating creation of report with a provided scan folder")

targetDirectorry = f'./{target}'
outputDirectory = f'{target}/output'
screenshotsDirectory = f'{target}/screenshots'
wordlistDirectory = f'{target}/wordlists'
credentailsDirectory = f'{target}/credentials'
reportsDirectory = f'{target}/reports'
extrasDirectory = f'{target}/extras'

outputNmapLocation = f'{outputDirectory}/outputNmap.xml'
outputNiktoLocation = f'{outputDirectory}/outputNikto.xml'
outputDirbLocation = f'{outputDirectory}/outputDirb.txt'
outputDirbLoggedLocation = f'{outputDirectory}/outputDirbLogged.txt'
docxReportLocation = f'{reportsDirectory}/{target}-scanReport.docx'
pdfReportLocation = f'{reportsDirectory}/'

wordlistUsernamesLocation = f'{wordlistDirectory}/wordlistUsernames.txt'
wordlistPasswordsLocation = f'{wordlistDirectory}/wordlistPasswords.txt'
credentialsTXTLocation = f'{credentailsDirectory}/credentials.txt'
credentialsJSONLocation = f'{credentailsDirectory}/credentials.json'

loginFailLocation = f'{extrasDirectory}/loginfail.txt'

def main():
    # Logged variable set to TRUE when succesfull hyndra
    logged = False
    # Is401 variable set to TRUE when there is a page that returns code 401 - Unauthorized Access
    is401 = False

    # Create directories
    if os.path.isdir(targetDirectorry):
        shutil.rmtree(targetDirectorry)
    os.mkdir(targetDirectorry)
    os.mkdir(outputDirectory)

    # Run the scans and get the output
    runNmap(target)
    runNikto(target)
    runDirb(target, logged)

    #Get addresses
    addressesOK = getOKAddr()
    addressesUNAUTH = getUNAUTHAddr()

    # Runs if user used the -L switch to turn the login script on
    if screenshotScript:
        grabScrenshot(addressesOK)

    nikto_vulns = []
    # Adds content to nikto vulns list
    extractInfoNikto(outputNiktoLocation, nikto_vulns)

    # Runs if user used the -L switch to turn the login script on
    if loginScript:
        autoLogin(addressesUNAUTH, addressesOK, nikto_vulns, is401, logged)

    createReport(logged)

    # Runs if user used the -P switch to turn output to pdf on
    if pdf:
        convertDocxToPdf()

# Function for running nmap
def runNmap(target):
    nmap_command = [f'sudo nmap -v -p- -A -oX {outputNmapLocation} {target}']
    subprocess.run(nmap_command, shell=True)

# Function for running nikto
def runNikto(target):
    nikto_command = [f'nikto -h {target} -output {outputNiktoLocation}']
    subprocess.run(nikto_command, shell=True)

# Function for running dirb
def runDirb(target, logged):
    # If succesfully logged in with credentials
    if logged:
        f = open(credentialsTXTLocation, "r")
        credentials = f.read()

        # Launch dirb with credentials provided 
        dirbuster_command = [f'dirb http://{target} -u {credentials} -o {outputDirbLoggedLocation}']
        subprocess.run(dirbuster_command, shell = True)
    else:
        # Launch dirb without credentials
        dirbuster_command = [f'dirb http://{target} -o {outputDirbLocation}']
        subprocess.run(dirbuster_command, shell=True)

# Get addresses from dirb output that returned 200. Addresses are displayed as http://address
def getOKAddr():
    command = [f"cat {outputDirbLocation} | grep CODE:200 | cut -d ' ' -f 2"]
    result = subprocess.run(command, capture_output=True, text=True, shell=True)
    # Split the output into lines
    return result.stdout.splitlines()

# Get addresses from dirb output that returned 401. Addresses are displayed as http://address
def getUNAUTHAddr():
    command = [f"cat {outputDirbLocation} | grep CODE:401 | cut -d ' ' -f 2"]
    result = subprocess.run(command, capture_output=True, text=True, shell=True)
    # Split the output into lines
    return result.stdout.splitlines()

# Function used for extracting information from nmap xml file
def extractInfoNmap(xml_file):
    # Create xml tree
    tree = ET.parse(xml_file)

    # Get root element
    root = tree.getroot()

    # Extract scan information
    scan_args = root.get('args')
    scan_start = root.get('startstr')
    scan_version = root.get('version')

    # Extract host information
    host = root.find('host')
    host_address = host.find('address').get('addr')

    # Extract open ports information
    open_ports = []
    ports = host.find('ports')
    for port in ports.findall('port'):
        ran_scripts = []
        port_id = port.get('portid')
        port_protocol = port.get('protocol')
        service = port.find('service')
        service_name = service.get('name')
        service_product = service.get('product')
        service_version = service.get('version')
        for script in port.findall('script'):
            script_id = script.get('id')
            script_output = script.get('output')
            ran_scripts.append({
                'script_id' : script_id,
                'script_output' : script_output
            })
        open_ports.append({
            'port_id': port_id,
            'protocol': port_protocol,
            'service_name': service_name,
            'service_product': service_product,
            'service_version': service_version,
            'ran_scripts': ran_scripts
        })

    # Extracted info
    output_string = f"Scan Information:\n" \
               f"======================================================================\n" \
               f"Scan Arguments: {scan_args}\n" \
               f"Scan Start Time: {scan_start}\n" \
               f"Scan Version: {scan_version}\n" \
               f"======================================================================\n" \
               f"Host Information:\n" \
               f"Host Address: {host_address}\n" \
               f"======================================================================\n" \
               f"Open Ports:\n" \
               f"----------------------------------------------------------------------------------------------------------------------\n"
    
    for port in open_ports:
        output_string += f"Port: {port['port_id']}\n" \
                        f"Protocol: {port['protocol']}\n" \
                        f"Service Name: {port['service_name']}\n" \
                        f"Service Product: {port['service_product']}\n" \
                        f"Service Version: {port['service_version']}\n" \
                        f"Scripts ran:\n"
        
        for script in port['ran_scripts']:
            output_string += f"-> {script['script_id']}: {script['script_output']}\n"

        output_string += f"----------------------------------------------------------------------------------------------------------------------\n"

    output_string += "======================================================================"
    return output_string

# Function used for extracting information from nikto xml file
def extractInfoNikto(xml_file, niktoVulns):
    # Create xml tree
    tree = ET.parse(xml_file)

    # Get root element
    root = tree.getroot()
    
    # Find niktoscan element inside the root
    niktoscan = root.find('niktoscan')

    # Find scan details element inside nikoscan
    scandetails = niktoscan.find('scandetails')

    # Extract target information
    target_ip = scandetails.get('targetip')
    target_hostname = scandetails.get('targethostname')
    target_port = scandetails.get('targetport')
    starttime = scandetails.get('starttime')

    # Extract All items
    for item in scandetails.findall('item'):
        method = item.get('method')
        description = item.find('description').text.strip()
        referencesField = item.find('references')
        if referencesField.text:
            references = referencesField.text.strip()
        else:
            references = 'Reference not available'
        uri = item.find('uri').text.strip()
        niktoVulns.append({
            'method' : method,
            'description' : description,
            'uri' : uri,
            'references' : references
        })

    # Extracted info
    output_string = f"Scan Information:\n" \
               f"======================================================================\n" \
               f"Target IP: {target_ip}\n" \
               f"Target hostname: {target_hostname}\n" \
               f"Target port: {target_port}\n" \
               f"Scan start time: {starttime}\n" \
               f"======================================================================\n" \
               f"Items:\n" \
               f"----------------------------------------------------------------------------------------------------------------------\n"

    for niktoVuln in niktoVulns:
        output_string += f"Method: {niktoVuln['method']}\n" \
                        f"Description: {niktoVuln['description']}\n" \
                        f"URI: {niktoVuln['uri']}\n" \
                        f"Reference: {niktoVuln['references']}\n" \
                        f"----------------------------------------------------------------------------------------------------------------------\n"

    output_string += "======================================================================"

    return output_string

# Function for taking screenshots of all pages that returned 200
def grabScrenshot(addressesOK):
    # Create directory
    os.mkdir(screenshotsDirectory)

    # Iterate over each address
    for addressOK in addressesOK:
        # Create driver
        driver = webdriver.Firefox()

        # Open address
        driver.get(addressOK)

        sleep(1)
        # Remove ":" and "/" from address using replace()
        file_name = addressOK.replace(":", ".").replace("/", ".")

        # Take a screenshot and save it
        driver.get_screenshot_as_file(f'{screenshotsDirectory}/{file_name}.png')

        # Exit driver
        driver.quit()

# Function used to find if nikto found a login section
def checkIfLoginNikto(nikto_vulns):
    # Checks if login page exists
    checkMessage = 'Admin login page/section found.'

    for vuln in nikto_vulns:
        # If nikto discovered login page
        if checkMessage in vuln['description']:
            return vuln['uri']

# Functions used for retriving text and comments from the pages
def getComments(addressesOK):
    print("Getting comments")

    # Create wordlist usernames file
    wordlistUsernamesFile = open(wordlistUsernamesLocation, "w+")

    # Loop for each address that returned 200
    for addressOK in addressesOK:
        # Gate all the html code for the address
        page_html = requests.get(addressOK).text

        # Parse html using beautifulsoup
        soup = BeautifulSoup(page_html, "lxml")

        # Find all comments
        comments = soup.find_all(string=lambda text: isinstance(text, Comment))

        # Loop for each comment that was found
        for comment in comments:
            # Split aech comment onto a list of words
            words = comment.split()

            # For each word in the list of words
            for word in words:
                # Write to username wordlist file
                wordlistUsernamesFile.write(word + '\n')

        # Get all the text
        text_content = soup.get_text()

        # Split the text onto a list of words
        words = text_content.split()

        # For each word in the list of words
        for word in words:
            # Write to username wordlist file
            wordlistUsernamesFile.write(word + '\n')
    
    # Sets the position in the file to the start
    wordlistUsernamesFile.seek(0)

    # Read the contents of the wordlist file which contains special characters
    wordlistNotClean = wordlistUsernamesFile.read()

    # The regex [^a-zA-Z0-9\s] will match any character that is not alphanumeric, a space, or a newline.
    disallowed_chars_regex = r'[^a-zA-Z0-9\s]'

    # Remove disallowed characters from the file contents
    wordlistClean = re.sub(disallowed_chars_regex, '', wordlistNotClean)

    # Sets the position in the file to the start
    wordlistUsernamesFile.seek(0)

    # Empties the file
    wordlistUsernamesFile.truncate(0)

    # Write the clean wordlist without special characters into the wordlist usernames file
    wordlistUsernamesFile.write(wordlistClean)

    # Closes the file
    wordlistUsernamesFile.close()

    # Opens rockyou.txt and reads it contents
    rockyou = open("rockyou.txt", "r", encoding="utf-8", errors="ignore")
    rockyou_contents = rockyou.read()
    rockyou.close()

    # Creates wordlist passwords file and opens it for writing
    wordlistPasswordsFile = open(wordlistPasswordsLocation, "w")

    # Prites the clean wordlist onto the opened file
    wordlistPasswordsFile.write(wordlistClean)

    # Prints the rockyou wordlist onto the opened file
    # This is done so only one file has the contents of rockyou, because bruteforcing both username and password will be too long
    wordlistPasswordsFile.write(rockyou_contents)
    wordlistPasswordsFile.close

# Function used getting all the login fields
def getLoginFields(target, uri):
    print("Getting Login Fields")

    # Gate all the html code for the address
    page_html = requests.get(f"http://{target}{uri}").text

    # Parse html using beautifulsoup
    soup = BeautifulSoup(page_html, "lxml")

    # Find login form
    login_form = soup.find("form")

    # Retrived method used for logging
    method = login_form.get("method")

    # Find input fields in login form
    fields = login_form.find_all("input")

    # List for contents of input fields
    fieldContents = []

    # Loop for each field that was found
    for field in fields:
        # Get required values
        name_field = field.get("name")
        value_field = field.get("value")
        type_field = field.get("type")

        # Append a dictionary to the list
        fieldContents.append({
            'name_field': name_field,
            'value_field': value_field,
            'type': type_field
        })

    return fieldContents, method

# Function used for getting the message displayed on unsuccessful login
def getLoginFail(target, uri, fieldContents):
    print("Getting Login Fail Message")

    # Take the contents of the fields
    for fieldDictionary in fieldContents:
            if fieldDictionary['type'] == "text":
                username_name = fieldDictionary['name_field']
            elif fieldDictionary['type'] == "password":
                password_name = fieldDictionary['name_field']
            elif fieldDictionary['type'] == "submit":
                submit_name = fieldDictionary['name_field']
                submit_value = fieldDictionary['value_field']
            else:
                print("Automatic retrivel of login fail message unsucessfull. Create hydra command manually")

    login_url = f"http://{target}{uri}"
    usernameContents = ""
    passwordContents = ""
    submitContent = submit_value

    # Send a POST request with login credentials
    payload = {
        username_name: usernameContents,
        password_name: passwordContents,
        submit_name : submitContent
    }

    # Get html of login page before and after attempting a login
    request_html = requests.get(login_url).text
    response_html = requests.post(login_url, data=payload).text
    # Parse the HTML responses
    soup1 = BeautifulSoup(request_html, 'lxml')
    soup2 = BeautifulSoup(response_html, 'lxml')
    # Convert the parsed HTML back to strings for comparison
    html_str1 = str(soup1)
    html_str2 = str(soup2)

    # Compare the html of both
    differ = difflib.Differ()
    differences = list(differ.compare(html_str1.splitlines(), html_str2.splitlines()))
    f = open(loginFailLocation, "w+")
    for diff in differences:
        if diff.startswith('+ '):
            f.write(diff[2:])

    f.seek(0)
    login_fail_contents = f.read()
    f.close()

    # Find the index of the first ">" character
    start_index = login_fail_contents.find(">") + 1
    # Find the index of the second "<" character, starting from the position after the first ">"
    end_index = login_fail_contents.find("<", start_index)
    # Extract the desired substring
    extracted_text = login_fail_contents[start_index:end_index].strip()

    return extracted_text

# Function used creating hydra command
def hydraBuilder(target, uri, is401, fieldContents, method):
    print("Creating hydra command")
    
    # Checks if it nees to build a command for a 401 page
    if is401:
        hydra_command = f"hydra -f -u -o {credentialsJSONLocation} -b json -L {wordlistUsernamesLocation} -P {wordlistPasswordsLocation} {target} http-get {uri}"
    else:
        # Gets messaged displayed on failed login
        login_fail = getLoginFail(target, uri, fieldContents)
        if method == "post":
            method = "http-post-form"
        else:
            print("Error getting method for hydra command. Try using hydra manually using wordlists found in wordlists directory")

        username_field = None
        password_field = None
        extra_field = None

        # Loop for each field that was found from getLoginFields
        for fieldDictonary in fieldContents:
            if fieldDictonary['type'] == "text":
                username_field = fieldDictonary['name_field']
            elif fieldDictonary['type'] == "password":
                password_field = fieldDictonary['name_field']
            elif fieldDictonary['type'] == "submit":
                extra_field = {
                    'name' : fieldDictonary['name_field'],
                    'value' : fieldDictonary['value_field']
                }
            else:
                print("Unable to setup hydra command. Use manual hyndra")

        form = f"{uri}:{username_field}=^USER^&{password_field}=^PASS^&{extra_field['name']}={extra_field['value']}:{login_fail}"
        print(f"FORM IS: {form}")
        hydra_command = f"hydra -f -u -o {credentialsJSONLocation} -b json -L {wordlistUsernamesLocation} -P {wordlistPasswordsLocation} {target} {method} '{form}'"
        
    subprocess.run(hydra_command, shell=True)   

# Function used for auto login into pages
def autoLogin(addressesUNAUTH, addressesOK, nikto_vulns, is401, logged):
    # Create directories
    os.mkdir(wordlistDirectory)
    os.mkdir(credentailsDirectory)
    os.mkdir(extrasDirectory)

    # Check if nikto discover a login page and set uri to uri of login page
    uri = checkIfLoginNikto(nikto_vulns)

    # If login page exist uri is set to a value and this executes
    if uri:
        print("Nikto login found. Creating wordlist")
        print("Login page uri:", uri)

        # Gets all the text from the pages that returned 200
        getComments(addressesOK)

        # Gets the information about the login field and the method used for login
        fieldContents, method = getLoginFields(target, uri)

        # Used to build hydra command
        hydraBuilder(target, uri, is401, fieldContents, method)

    # If there is pages that returned 401
    elif addressesUNAUTH:
        is401 = True
        print("401 page exists. Creating wordlist from page contents")
        print("401 page Address:", addressesUNAUTH)

        # Gets all the text from the pages that returned 200
        getComments(addressesOK)

        # Gets the url for page that returned 401 in the format http://10.10.10.10/uri
        url = addressesUNAUTH[0]

        # Splits url at every '/'
        parts = url.split("/")

        # Gives everything to the right of the 3rd '/' which is the uri
        result = "/".join(parts[3:])

        # Sets the uri to be /uri
        uri = "/" + result

        # Method/login field are not needed for http-get login so they are set to none
        fieldContents = None
        method = None
        hydraBuilder(target, uri, is401, fieldContents, method)
        
    with open(credentialsJSONLocation, 'r') as json_file:
        data = json.load(json_file)

    # Extract login and password and from json file if it contains a sucessfull result
    if len(data['results']) > 0:

        # Extract username and password from json file
        username,password = extractCredentials()

        # Set logged to true to indicate login credentials were found
        logged = True

        # Take screenshots of pages after logging in
        loginPageAddress = target + uri
        
        if screenshotScript:
            grabScrenshotLogged(loginPageAddress, addressesUNAUTH, username, password)

        # Run dirb again with logged set to true which makes it use credentials
        runDirb(target, logged)

# Function used for extracting credentails that were found from credentials.json file created by hydra
def extractCredentials():
    #readJsonfile if it exist
    print(f"Reading credentails from {credentialsJSONLocation}")
    with open(credentialsJSONLocation, 'r') as json_file:
        data = json.load(json_file)

    login = data['results'][0]['login']
    password = data['results'][0]['password']
    # Save login and password to a file
    with open(credentialsTXTLocation, 'w') as file:
        file.write(f"{login}:{password}")
    print("Login and password saved successfully.")
    return login,password

# Function for taking screenshots of pages that require logging in
def grabScrenshotLogged(loginPageAddress, addressesUNAUTH, username, password):
    # Check if there was address that returned 401
    if addressesUNAUTH:
        # Iterate over each address
        for addressUNAUTH in addressesUNAUTH:
            # Create driver
            driver = webdriver.Firefox()

            # Get everything after http://
            split = addressUNAUTH.split("//")[1]

            # Open address usng username and password for authentication
            driver.get(f"http://{username}:{password}@{split}")

            sleep(1)

            # Remove ":" and "/" from address using replace()
            file_name = addressUNAUTH.replace(":", ".").replace("/", ".")

            # Take a screenshot and save it
            driver.get_screenshot_as_file(f'{screenshotsDirectory}/{file_name}.png')

            # Exit driver
            driver.quit()
    # Runs if there was no 401 pages, but there was a login page
    else:
        print("Does not have 401 pages")

        # Create driver
        driver = webdriver.Firefox()

        # Open address
        driver.get(f"http://{loginPageAddress}")

        # Gate all the html code for the address
        page_html = requests.get(f"http://{loginPageAddress}").text

        # Parse html using BeautifulSoup
        soup = BeautifulSoup(page_html, "lxml")

        # Find login form
        login_form = soup.find("form")

        # Find input fields in login form
        fields = login_form.find_all("input")

        # List for contents of input fields
        fieldContents = []

        # Loop for each field that was found
        for field in fields:
            # Get required values
            name_field = field.get("name")
            value_field = field.get("value")
            type_field = field.get("type")

            # Append a dictionary to the list
            fieldContents.append({
                'name_field': name_field,
                'value_field': value_field,
                'type': type_field
            })

        # Loop for each dictionary in the fieldContents lists
        for fieldDictionary in fieldContents:
            if fieldDictionary['type'] == "text":
                username_name = fieldDictionary['name_field']
            elif fieldDictionary['type'] == "password":
                password_name = fieldDictionary['name_field']
            elif fieldDictionary['type'] == "submit":
                submit_name = fieldDictionary['name_field']
            else:
                print(f"Automatic login at http://{loginPageAddress} unsuccesful. Try logging in manually with credentials found in credentials/credentials.txt")

        # Locate username field using the name of the element for the driver to use at input
        usernameField = driver.find_element(By.NAME, username_name)

        # Locate password field using the name of the element for the driver to use at input
        passwordField = driver.find_element(By.NAME, password_name)

        # Send keys to fields
        usernameField.send_keys(username)
        passwordField.send_keys(password)

        # Locate submit button and click it
        driver.find_element(By.NAME, submit_name).click()

        sleep(1)

        # Remove ":" and "/" from address using replace()
        file_name = loginPageAddress.replace("/", ".")

        # Take a screenshot and save it
        driver.get_screenshot_as_file(f'{screenshotsDirectory}/{file_name}.png')
    
        # Exit driver
        driver.quit()

# Function for saving output to word file
def saveOutputWord(filename, tool_name, output):
    # Opening docx file
    document = Document(filename)

    # Create a table and add it to document
    table = document.add_table(rows=2, cols=1)

    # Get cell on first row
    toolNameCell = table.cell(0, 0)
    toolNameCell.text = tool_name

    # Styling
    toolNameCell_paragraph = toolNameCell.paragraphs[0]
    toolNameCell_paragraph.style = document.styles['Heading 1']
    toolNameCell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Get cell on second row
    contentCell = table.cell(1, 0)
    contentCell.text = output

    document.save(filename)
    print(f"{tool_name} saved to {filename}")

# Function adding created images to word file
def addImageWord(filename):
    # Opening docx file
    document = Document(filename)
    
    # Create heading
    document.add_heading("Captured Screenshots", level=0)

    # iterate over screenshots in screenshots directory
    for screenshot in os.listdir(screenshotsDirectory):

        # Create path to screenshots by joining hte screenshots directory location and the screenshot name
        screenshotFile = os.path.join(screenshotsDirectory, screenshot)

        # checking if file exists at location
        if os.path.isfile(screenshotFile):
            # Create a table and add it to document
            table = document.add_table(rows=2, cols=1)

            # Get cell on first row
            nameCell = table.cell(0, 0)
            nameCell.text = f"Screenshot: {screenshotFile}"

            # Style
            nameCell_paragraph = nameCell.paragraphs[0]
            nameCell_paragraph.style = document.styles['Heading 1']
            nameCell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Get cell on second row
            imageCell = table.cell(1, 0)
            imageCell_paragraph = imageCell.paragraphs[0]

            # Adding picture to paragraph in cell
            imageCell_paragraph.add_run().add_picture(screenshotFile, width=Cm(15), height=None)
    
    document.save(filename)
    print(f"Added picture to {filename}")

# Function for adding created wordlist to word file
def addWordlistWord(filename):
    # Opening docx file
    document = Document(filename)
    
    # Creating a heading
    document.add_heading("Created wordlist", level=0)

    # Opening wordlist
    wordlistUsernamesFile = open(wordlistUsernamesLocation)

    # Reading wordlist
    content = wordlistUsernamesFile.read()

    # Adding contents of wordlist file to document
    document.add_paragraph(content)
    
    document.save(filename)
    print(f"Added comments to {filename}")

# Function for adding created credentials to word file
def addCredentialsWord(filename):
    # Opening docx file
    document = Document(filename)

    # Creating a heading
    document.add_heading("Discovered credentials", level=0)

    # If credentials worked for logging in
    if os.path.exists(credentialsTXTLocation):
        # Open credentials file
        credentialsTXTFile = open(credentialsTXTLocation)
        
        # Read credentials files
        content = credentialsTXTFile.read()

        # Adding contents of credentials file to document
        document.add_paragraph(content)
    else:
        document.add_paragraph("No credentials were found")

    print(f"Added credentials to {filename}")

    document.save(filename)

# Function for conveting docx file to pdf
def convertDocxToPdf():
    # Read the DOCX file
    command = [f"libreoffice --headless --convert-to pdf {docxReportLocation} --outdir {pdfReportLocation}"]
    subprocess.run(command, shell=True)

# Function used for demonstrating report creation from provided output folder
def createReport(logged):
    if os.path.isdir(reportsDirectory):
        shutil.rmtree(reportsDirectory)
    os.mkdir(reportsDirectory)
    #Create document
    document = Document()
    # Add a title
    document.add_heading('Scans Output', level=0)
    document.save(docxReportLocation)

    # Saves nmap output in docx
    nmap_output = extractInfoNmap(outputNmapLocation)
    saveOutputWord(docxReportLocation, 'Nmap Scan Output', nmap_output)

    # Variables used for storing all vulnerabilities discovered by nikto
    nikto_vulns = []
    # Saves nikto output in docx
    nikto_output = extractInfoNikto(outputNiktoLocation, nikto_vulns)
    saveOutputWord(docxReportLocation, 'Nikto Scan Output', nikto_output)

    # Saves dirb output in docx
    dirb_output = open(outputDirbLocation, "r")
    saveOutputWord(docxReportLocation, 'Dirbuster Scan Output', dirb_output.read())
    dirb_output.close()

    # Saves dirb logged output in docx
    if logged:
        dirbLogged_output = open(outputDirbLoggedLocation, "r")
        saveOutputWord(docxReportLocation, 'Dirbuster Logged Scan Output', dirbLogged_output.read())
        dirbLogged_output.close()

    # Runs if user used the -S switch to turn the screenshot script on or the -0 switch to turn using output folder only on
    if screenshotScript or outputOnly:
        # Adds images to docx
        addImageWord(docxReportLocation)

    # Runs if user used the -L switch to turn the login script on or the -0 switch to turn using output folder only on
    if loginScript or outputOnly:
        # Adds wordlist to docx
        addWordlistWord(docxReportLocation)
        # Adds discovered credentials to docx
        addCredentialsWord(docxReportLocation)

# Runs if user used the -O switch to turn using output folder on
if outputOnly:
    createReport(True)
    convertDocxToPdf()
else:
    main()